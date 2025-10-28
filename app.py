# app.py
import streamlit as st
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches
import re
import os
import json
import zipfile
from typing import Dict
import io
import smtplib, ssl
from email.message import EmailMessage
import google.generativeai as genai

# --- CẤU HÌNH BẢO MẬT ---
try:
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
    APP_EMAIL      = st.secrets["APP_EMAIL"]
    APP_PASSWORD   = st.secrets["APP_PASSWORD"]
except Exception:
    st.warning("Không tìm thấy Streamlit Secrets. Đang sử dụng cấu hình local. Đừng quên thiết lập Secrets khi deploy!")
    GEMINI_API_KEY = "YOUR_GEMINI_API_KEY"
    APP_EMAIL      = "your_email@example.com"
    APP_PASSWORD   = "your_app_or_email_password"

# Cấu hình API key cho Gemini
try:
    genai.configure(api_key=GEMINI_API_KEY)
except Exception as e:
    st.error(f"Lỗi cấu hình Gemini API: {e}. Vui lòng kiểm tra lại API Key.")

#======================================================================
# PHẦN 0: HÀM KIỂM TRA BẮT BUỘC
#======================================================================

EMAIL_RE = re.compile(r"^[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}$")
REQUIRED_PLACEHOLDERS = ["TenCuocHop", "ThoiGianCuocHop", "DiaDiemCuocHop", "TenChuTri", "TenThuKy"]

def validate_inputs(
    template_option: str,
    transcript_file,
    template_file,
    meeting_name: str,
    meeting_time: str,
    meeting_location: str,
    meeting_chair: str,
    meeting_secretary: str,
    recipient_email: str,
    default_template_path: str = None
) -> bool:
    """
    Trả về True nếu hợp lệ; ngược lại hiển thị thông báo đỏ và trả về False.
    """
    missing = []

    # File bắt buộc
    if not transcript_file:
        missing.append("File transcript (.docx)")

    if template_option == "Template VPI":
        if default_template_path and not os.path.exists(default_template_path):
            missing.append(f"Template mặc định không tồn tại: {default_template_path}")
    elif template_option == "Template tùy chỉnh":
        if not template_file:
            missing.append("File template tùy chỉnh (.docx)")

    # Trường bắt buộc
    if not meeting_name:
        missing.append("Tên cuộc họp")
    if not meeting_time:
        missing.append("Thời gian cuộc họp")
    if not meeting_location:
        missing.append("Địa điểm cuộc họp")
    if not meeting_chair:
        missing.append("Tên chủ trì")
    if not meeting_secretary:
        missing.append("Tên thư ký")
    if not recipient_email:
        missing.append("Email nhận kết quả")
    elif not EMAIL_RE.match(recipient_email.strip()):
        missing.append("Email nhận kết quả (không hợp lệ)")

    if missing:
        st.error("❌ **Chưa hoàn thành thông tin**:\n\n- " + "\n- ".join(missing) +
                 "\n\nVui lòng bổ sung/đính kèm đầy đủ rồi bấm lại **Tạo biên bản**.")
        return False

    return True

#======================================================================
# PHẦN 1: HÀM XỬ LÝ (theo logic của .ipynb)
#======================================================================

# Regex y hệt notebook
COMMENT_RE     = re.compile(r"\{#.*?#\}")                 # 1-run
COMMENT_ALL_RE = re.compile(r"\{#.*?#\}", re.DOTALL)      # đa-run
BOLD_RE        = re.compile(r"\*\*(.*?)\*\*")             # **bold**
TOKEN_RE       = re.compile(r"\{\{([^{}]+)\}\}")          # {{Key}}

def _is_md_table(text: str) -> bool:
    lines = [l.strip() for l in (text or "").strip().splitlines() if l.strip()]
    return (
        len(lines) >= 2
        and "|" in lines[0]
        and set(lines[1].replace(" ", "").replace(":", "")) <= set("-|")
    )

def _parse_md_table(text: str):
    lines  = [l.strip() for l in (text or "").strip().splitlines() if l.strip()]
    header = [c.strip() for c in lines[0].split("|")]
    # bỏ cell rỗng do | đầu/cuối
    if header and header[0] == "":
        header = header[1:]
    if header and header[-1] == "":
        header = header[:-1]
    rows   = []
    for ln in lines[2:]:  # Skip header + separator
        cols = [c.strip() for c in ln.split("|")]
        if cols and cols[0] == "":
            cols = cols[1:]
        if cols and cols[-1] == "":
            cols = cols[:-1]
        if cols:
            while len(cols) > len(header):
                cols.pop()
            while len(cols) < len(header):
                cols.append("")
            rows.append(cols)
    return header, rows

def _insert_paragraph_after(anchor_para: Paragraph, style=None) -> Paragraph:
    """Chèn một đoạn (w:p) NGAY SAU anchor_para và trả về Paragraph mới."""
    new_p_ox = OxmlElement("w:p")
    anchor_para._p.addnext(new_p_ox)
    new_para = Paragraph(new_p_ox, anchor_para._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para

def add_formatted_text(paragraph: Paragraph, text: str, style_info=None):
    """
    Thêm text vào paragraph, hỗ trợ **bold** theo markdown và xuống dòng.
    style_info (nếu có) dùng để clone phông/chữ từ đoạn gốc.
    """
    parts   = BOLD_RE.split(text or "")
    is_bold = False
    for part in parts:
        if part == "":
            is_bold = not is_bold
            continue
        lines = part.split("\n")
        for i, ln in enumerate(lines):
            if i > 0:
                paragraph.add_run().add_break()
            if ln == "":
                continue
            run = paragraph.add_run(ln)
            if style_info:
                try:
                    f = run.font
                    if style_info.get("size"):
                        f.size = style_info["size"]
                    if style_info.get("name"):
                        f.name = style_info["name"]
                    if style_info.get("bold") is not None:
                        f.bold = style_info["bold"]
                    if style_info.get("italic") is not None:
                        f.italic = style_info["italic"]
                except Exception:
                    pass
            run.bold = run.bold or is_bold
        is_bold = not is_bold

def _concat_runs(paragraph: Paragraph):
    """Trả về (full_text, meta) với meta = [(run, start, end)]."""
    meta, pos, buf = [], 0, []
    for r in paragraph.runs:
        t = r.text or ""
        start, end = pos, pos + len(t)
        meta.append((r, start, end))
        buf.append(t)
        pos = end
    return "".join(buf), meta

def _insert_table_after(paragraph: Paragraph, header, rows, table_style="New Table"):
    """Chèn bảng sau một paragraph, từ header + rows (đã parse)."""
    if not header or not rows:
        return
    body = paragraph._parent  # có thể là Document hoặc Cell
    tbl  = body.add_table(rows=len(rows)+1, cols=len(header))
    try:
        tbl.style = table_style
    except Exception:
        pass
    # Header
    for i, h in enumerate(header):
        try:
            tbl.rows[0].cells[i].text = str(h)
        except Exception:
            pass
    # Rows
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, cell_val in enumerate(row):
            try:
                tbl.rows[r_idx].cells[c_idx].text = str(cell_val)
            except Exception:
                pass
    # Đặt bảng ngay sau đoạn anchor
    paragraph._p.addnext(tbl._tbl)

def extract_vars_and_desc(docx_file_or_buffer) -> Dict[str, str]:
    """Trích xuất placeholders {{Key}} {# mô tả #} từ .docx (đường dẫn hoặc buffer)."""
    xml_parts = []
    with zipfile.ZipFile(docx_file_or_buffer) as z:
        for name in z.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                xml_parts.append(z.read(name).decode("utf8"))
    all_xml = "\n".join(xml_parts)
    texts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", all_xml, flags=re.DOTALL)
    full_text = "".join(texts)
    pattern = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}\s*\{#\s*(.*?)\s*#\}", flags=re.DOTALL)
    return dict(pattern.findall(full_text))

def replace_in_paragraph(paragraph: Paragraph, data: Dict[str, str]):
    """
    Thay {{Key}} và xoá {#...#} đa-run; nếu value là:
    - bảng Markdown: chèn bảng ngay sau paragraph;
    - bullet (- / +): chèn các đoạn bullet ngay sau paragraph;
    - văn bản thường: thay trực tiếp giữ prefix/suffix giữa các run.
    """
    if not paragraph.runs:
        return

    full_text, meta = _concat_runs(paragraph)
    if not full_text:
        return

    # Gom comment + token thuộc data
    items = []
    for m in COMMENT_ALL_RE.finditer(full_text):
        items.append(("comment", m.start(), m.end(), None))
    for m in TOKEN_RE.finditer(full_text):
        key = (m.group(1) or "").strip()
        if key in data:
            items.append(("token", m.start(), m.end(), key))

    if not items:
        # fallback: xoá comment dạng 1-run
        for r in paragraph.runs:
            if r.text and COMMENT_RE.search(r.text):
                r.text = COMMENT_RE.sub("", r.text)
        return

    # Xử lý từ phải -> trái để không lệch chỉ số
    items.sort(key=lambda x: x[1], reverse=True)

    bullet_queue = []  # (text, style)
    table_queue  = []  # (header, rows)

    for item_type, start, end, key in items:
        # Tính vị trí run bao/chéo
        run_start_idx = next((i for i, (_, s, e) in enumerate(meta) if s <= start < e), None)
        run_end_idx   = next((i for i, (_, s, e) in enumerate(meta) if s <  end <= e), None)
        if run_start_idx is None or run_end_idx is None:
            continue

        run_start, s0, e0 = meta[run_start_idx]
        run_end,   s1, e1 = meta[run_end_idx]
        offset_start = start - s0
        offset_end   = end   - s1

        if item_type == "comment":
            # Xoá {#...#}
            if run_start_idx == run_end_idx:
                t = run_start.text or ""
                run_start.text = t[:offset_start] + t[offset_end:]
            else:
                run_start.text = (run_start.text or "")[:offset_start]
                for i in range(run_start_idx + 1, run_end_idx):
                    meta[i][0].text = ""
                run_end.text = (run_end.text or "")[offset_end:]
            continue

        # Token {{key}}
        value = data.get(key, "")

        # BẢNG MARKDOWN
        if isinstance(value, str) and _is_md_table(value):
            try:
                header, rows = _parse_md_table(value)
                table_queue.append((header, rows))
                # Xoá token khỏi đoạn
                if run_start_idx == run_end_idx:
                    t = run_start.text or ""
                    run_start.text = t[:offset_start] + t[offset_end:]
                else:
                    run_start.text = (run_start.text or "")[:offset_start]
                    for i in range(run_start_idx + 1, run_end_idx):
                        meta[i][0].text = ""
                    run_end.text = (run_end.text or "")[offset_end:]
                continue
            except Exception:
                # fallback về text thường
                value = str(value)

        # BULLET LIST (dòng bắt đầu bằng '-' hoặc '+')
        if isinstance(value, str) and any(line.strip().startswith(("-", "+")) for line in value.splitlines()):
            for line in value.splitlines():
                s = line.strip()
                if s.startswith("-"):
                    bullet_queue.append((s[1:].strip(), "List Bullet"))
                elif s.startswith("+"):
                    bullet_queue.append((s[1:].strip(), "List Bullet 2"))
            # Xoá token
            if run_start_idx == run_end_idx:
                t = run_start.text or ""
                run_start.text = t[:offset_start] + t[offset_end:]
            else:
                run_start.text = (run_start.text or "")[:offset_start]
                for i in range(run_start_idx + 1, run_end_idx):
                    meta[i][0].text = ""
                run_end.text = (run_end.text or "")[offset_end:]
            continue

        # VĂN BẢN THƯỜNG
        replacement_text = str(value)
        if run_start_idx == run_end_idx:
            t = run_start.text or ""
            run_start.text = t[:offset_start] + replacement_text + t[offset_end:]
        else:
            # clear phần giữa
            for i in range(run_start_idx + 1, run_end_idx):
                meta[i][0].text = ""
            # start run = prefix + replacement
            start_text = (run_start.text or "")[:offset_start]
            run_start.text = start_text + replacement_text
            # end run = suffix
            run_end.text = (run_end.text or "")[offset_end:]

    # Chèn bullet/bảng ngay sau paragraph
    if bullet_queue or table_queue:
        current_para = paragraph
        # bullets
        for text, style in bullet_queue:
            current_para = _insert_paragraph_after(current_para, style=style)
            add_formatted_text(current_para, text)
        # tables
        for header, rows in table_queue:
            try:
                _insert_table_after(current_para, header, rows)
            except Exception as e:
                print(f"Error inserting table: {e}")

def fill_template_to_buffer(template_file_or_path, data_input: Dict[str, str]):
    """Điền dữ liệu vào template và trả về BytesIO .docx (xử lý thân + bảng + header + footer)."""
    try:
        doc = Document(template_file_or_path)
    except Exception as e:
        st.error(f"Lỗi mở template: {e}")
        return None

    # Body
    for i, paragraph in enumerate(doc.paragraphs):
        try:
            replace_in_paragraph(paragraph, data_input)
        except Exception as e:
            print(f"Error processing paragraph {i}: {e}")

    # Tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, paragraph in enumerate(cell.paragraphs):
                    try:
                        replace_in_paragraph(paragraph, data_input)
                    except Exception as e:
                        print(f"Error processing table {table_idx}, row {row_idx}, cell {cell_idx}, paragraph {para_idx}: {e}")

    # Headers & Footers
    for section in doc.sections:
        # Header
        if section.header:
            for paragraph in section.header.paragraphs:
                try:
                    replace_in_paragraph(paragraph, data_input)
                except Exception as e:
                    print(f"Error processing header paragraph: {e}")
        # Footer
        if section.footer:
            for paragraph in section.footer.paragraphs:
                try:
                    replace_in_paragraph(paragraph, data_input)
                except Exception as e:
                    print(f"Error processing footer paragraph: {e}")

    # Lưu vào buffer
    try:
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    except Exception as e:
        st.error(f"Đã xảy ra lỗi khi tạo file Word: {e}")
        return None

def call_gemini_model(transcript_content, placeholders):
    """Gửi yêu cầu đến Gemini và nhận về kết quả JSON (giữ prompt như notebook)."""
    model = genai.GenerativeModel("gemini-2.5-pro")
    Prompt_word = """
# Vai trò
Bạn là một trợ lý AI chuyên nghiệp, có nhiệm vụ trích xuất thông tin quan trọng từ bản ghi cuộc họp để tạo ra nội dung cho biên bản cuộc họp, đảm bảo tính chính xác và trình bày chuyên nghiệp.

# Đầu vào
1.  **Bản ghi cuộc họp (transcript):** `{0}`
2.  **Danh sách các trường thông tin cần trích xuất (placeholders):** `{1}` (Đây là một đối tượng/dictionary nơi mỗi key là tên trường cần trích xuất và value là mô tả/yêu cầu định dạng cho trường đó).

# Nhiệm vụ
1.  **Phân tích kỹ lưỡng:** Đọc và hiểu toàn bộ nội dung bản ghi cuộc họp.
2.  **Xác định và Trích xuất:** Với **từng** trường thông tin (key) trong danh sách `placeholders`:
    *   Tìm (các) phần nội dung tương ứng trong bản ghi.
    *   Trích xuất thông tin một cách **chi tiết, đầy đủ ý, và chính xác tuyệt đối** về mặt ngữ nghĩa so với bản ghi gốc.
    *   **Trường hợp không có thông tin:** Nếu không tìm thấy thông tin rõ ràng cho một trường cụ thể trong bản ghi, hãy ghi nhận là "Chưa có thông tin".
3.  **Định dạng và Diễn đạt:**
    *   **Luôn trả về bằng tiếng Việt.**
    *   Sử dụng **văn phong trang trọng, lịch sự, chuyên nghiệp**, phù hợp với tiêu chuẩn của một biên bản cuộc họp chính thức.
    *   Diễn đạt thành **câu văn hoàn chỉnh, mạch lạc, đúng ngữ pháp và chính tả tiếng Việt**. Tổng hợp các ý rời rạc hoặc văn nói thành cấu trúc văn viết chuẩn mực.
    *   Đảm bảo mỗi thông tin trích xuất đều **rõ ràng, súc tích và có ý nghĩa**.
    *   **Quan trọng:** Áp dụng **đúng định dạng trình bày** (ví dụ: bullet cấp 1, bullet cấp 2, bảng Markdown, đoạn văn...) **theo yêu cầu được chỉ định trong phần mô tả (value) của placeholder tương ứng**.
4.  **Tạo đối tượng JSON:** Tập hợp tất cả thông tin đã trích xuất và định dạng vào một đối tượng JSON duy nhất, tuân thủ nghiêm ngặt các quy tắc xuất kết quả.

# Quy tắc xuất kết quả (Quan trọng - Tuân thủ nghiêm ngặt)
1.  **Khóa (keys) của JSON:**
    *   Phải **trùng khớp 100%** với từng phần tử (key) trong danh sách `placeholders`.
    *   Giữ nguyên mọi ký tự: dấu, dấu câu, khoảng trắng, chữ hoa/thường.
    *   **Tuyệt đối không:** chuyển sang không dấu, snake_case, camelCase, viết tắt, hoặc thay đổi tên khóa.
2.  **Cấu trúc JSON:**
    *   Chỉ xuất các cặp key-value tương ứng với `placeholders`.
    *   **Không** thêm khóa mới, **không** bớt khóa, **không** lồng ghép cấu trúc khác.
3.  **Giá trị (values) của JSON:**
    *   **Tuân thủ Yêu cầu Định dạng từ Placeholder:** **Đây là điểm cực kỳ quan trọng.** Đối với **mỗi** trường thông tin (key) trong JSON, bạn phải **đọc kỹ yêu cầu định dạng được nêu trong phần mô tả (value) của placeholder tương ứng** trong danh sách `placeholders`. **Áp dụng chính xác** định dạng đó cho chuỗi giá trị (value) của trường đó.
        *   Ví dụ: Nếu placeholder có yêu cấu trình bày theo bullet cấp 2 thì giá trị value trong Json phải bắt đầu mỗi dòng bằng '+'; hoặc nếu placeholder yêu cầu trình bày là dạng bảng thì giá trị key trong Json phải bắt buộc là dạng bảng markdown.
    *   **Nội dung:** Phải là kết quả đã được xử lý theo **Mục 3 (Định dạng và Diễn đạt)** ở phần Nhiệm vụ, đồng thời được **trình bày một cách rõ ràng, có cấu trúc chặt chẽ, và chuyên nghiệp** theo đúng yêu cầu định dạng từ placeholder.
    *   **Kiểu dữ liệu:** Tất cả giá trị (values) trong JSON phải là kiểu **chuỗi (string)**. **Tuyệt đối không sử dụng kiểu mảng (array) hoặc các kiểu dữ liệu khác.**
    *   **Xử lý trường hợp không có thông tin:** Nếu không tìm thấy thông tin cho một trường cụ thể trong bản ghi, giá trị tương ứng trong JSON phải là chuỗi: `Chưa có thông tin`.
    *   **Hướng dẫn Định dạng Bullet (KHI được yêu cầu trong Placeholder):** Mục tiêu là tạo ra văn bản có cấu trúc, dễ đọc và chuyên nghiệp. **Toàn bộ cấu trúc này phải được thể hiện bên trong chuỗi giá trị.**
        *   **Bullet cấp 1 (Thường dùng cho mục chính):** Bắt đầu dòng bằng dấu gạch ngang theo sau là một khoảng trắng (`- `) cho mỗi ý chính.
        *   **Bullet cấp 2 (Thường dùng cho ý phụ, chi tiết):** Bắt đầu dòng bằng dấu cộng theo sau là một khoảng trắng (`+ `) cho mỗi ý phụ. Nên thụt lề đầu dòng cho các mục cấp 2 (ví dụ: thêm 2 hoặc 4 dấu cách trước dấu `+ `) để phân biệt rõ ràng với cấp 1.
        *   **Trình bày dòng:** Mỗi mục bullet (cả `- ` và `+ `) phải nằm trên một dòng riêng biệt trong chuỗi kết quả. AI cần đảm bảo việc xuống dòng phù hợp giữa các mục bullet để tạo cấu trúc danh sách rõ ràng khi chuỗi được hiển thị.
        *   **Đặc biệt với Công việc cần làm (Action Items) (NẾU placeholder yêu cầu cấu trúc này):** Cấu trúc rõ ràng thông tin cho từng mục, ví dụ sử dụng bullet cấp 1 (`- `) cho mỗi công việc và bullet cấp 2 (`+ `) thụt lề cho các chi tiết:
            - [Nội dung công việc cụ thể 1]
              + Người phụ trách: [Tên người/Bộ phận]
              + Hạn chót: [Ngày/Thời hạn cụ thể]
            - [Nội dung công việc cụ thể 2]
              + Người phụ trách: [Tên người/Bộ phận]
              + Hạn chót: [Ngày/Thời hạn cụ thể]
        *   **Tính nhất quán:** Áp dụng định dạng (bullet, bảng, đoạn văn...) một cách nhất quán theo đúng yêu cầu của từng placeholder.
4.  **Định dạng đầu ra:**
    *   **Không** bao gồm bất kỳ chú thích, giải thích, lời dẫn nào bên ngoài đối tượng JSON (ví dụ: không có `Đây là kết quả:` hay ```json ... ```).
    *   Toàn bộ kết quả trả về phải là **một chuỗi JSON hợp lệ và duy nhất**.
    """
    prompt = Prompt_word.format(transcript_content, placeholders)
    try:
        response = model.generate_content(
            contents=prompt,
            generation_config={"response_mime_type": "application/json"}
        )
        if response and hasattr(response, "text"):
            raw = response.text.strip()
            if raw.startswith("```"):
                raw = raw.split("```")[1].strip("json\n")
            return json.loads(raw)
        else:
            st.error("Phản hồi từ Gemini API bị thiếu hoặc không hợp lệ.")
            return None
    except Exception as e:
        st.error(f"Lỗi khi gọi Gemini API: {e}")
        return None

def send_email_with_attachment(recipient_email, attachment_buffer, filename="BBCH.docx"):
    """Gửi email với file đính kèm từ buffer."""
    SMTP_SERVER = "smtp.office365.com"
    SMTP_PORT = 587

    msg = EmailMessage()
    msg["Subject"] = "Biên bản cuộc họp đã được tạo tự động"
    msg["From"] = APP_EMAIL
    msg["To"] = recipient_email
    msg.set_content(
        "Chào bạn,\n\nBiên bản cuộc họp đã được tạo thành công.\nVui lòng xem trong file đính kèm.\n\nTrân trọng,\nCông cụ tạo biên bản tự động."
    )
    msg.add_attachment(
        attachment_buffer.getvalue(),
        maintype="application",
        subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )

    try:
        ctx = ssl.create_default_context()
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as s:
            s.starttls(context=ctx)
            s.login(APP_EMAIL, APP_PASSWORD)
            s.send_message(msg)
        return True
    except Exception as e:
        st.error(f"Lỗi khi gửi email: {e}. Vui lòng kiểm tra lại cấu hình email và mật khẩu ứng dụng.")
        return False

#======================================================================
# PHẦN 2: GIAO DIỆN STREAMLIT
#======================================================================

st.set_page_config(layout="wide", page_title="Công cụ tạo Biên bản cuộc họp")
st.title("🛠️ Công cụ tạo biên bản cuộc họp tự động")

with st.sidebar:
    st.info("📝 **Hướng dẫn sử dụng**")
    st.markdown("""
    1.  **Tải file transcript:** Tải lên file `.docx` chứa nội dung cuộc họp.
    2.  **Chọn Template:**
        * Sử dụng mẫu có sẵn bằng cách chọn "Template VPI".
        * Hoặc "Template tùy chỉnh" và tải file của bạn lên.
    3.  **Điền thông tin:** Nhập các thông tin cơ bản của cuộc họp.
    4.  **Nhập email:** Điền địa chỉ email bạn muốn nhận kết quả.
    5.  **Chạy:** Nhấn nút 'Tạo biên bản'.
    """)

    
    st.info("📝 **Hướng dẫn tạo template**")
    st.markdown("""
📂 File nhận đầu vào là file có đuôi `.docx`
Khi tạo template cho biên bản cuộc họp, bạn cần mô tả rõ từng biến để đảm bảo hệ thống hiểu đúng và điền thông tin chính xác. Mỗi biến cần tuân thủ cấu trúc sau: 
{{Ten_bien}}{# Mo_ta_chi_tiet #}
🔍 Trong đó:
- ✅ {{Ten_bien}}:
- Tên biến **viết bằng tiếng Anh hoặc tiếng Việt không dấu**.
- **Không sử dụng dấu cách**. Nếu cần phân tách các từ, sử dụng **dấu gạch dưới (_)**
- Dấu ngoặc nhọn kép ({{ và }}) phải **gắn liền với tên biến**, **không có khoảng trắng**.
- Ví dụ hợp lệ: {{Thanh_phan_tham_du}}
- ✅ {# Mo_ta_chi_tiet #}:
- Mở đầu bằng dấu {#, tiếp theo là nội dung mô tả, và kết thúc bằng dấu #}.
- Nội dung mô tả phải nêu rõ:
  - **Thông tin cần điền** vào biến là gì (dữ liệu nội dung).
  - **Yêu cầu trình bày** dữ liệu như thế nào (ví dụ: dạng bảng, dạng bullet,...).
  - **Cấu trúc trình bày:** chỉ hỗ trợ **hai cấp trình bày**:
    - **Bullet cấp 1**: dùng cho ý chính
    - **Bullet cấp 2**: dùng cho các ý nhỏ bổ sung dưới từng ý chính.
🧾 Ví dụ cụ thể:
{{Thanh_phan_tham_du}}{#Danh sách người tham gia cuộc họp, trình bày ở dạng bullet point. Ưu tiên sắp xếp từ lãnh đạo cấp cao, lãnh đạo bộ phận đến chuyên viên. Chỉ sử dụng tối đa 2 cấp trình bày: bullet 1 là tên từng người, bullet 2 là chức vụ hoặc vai trò nếu có.#}

- **🎨 Tạo định dạng hiển thị cho các bullet:**
- 📍 Đối với bullet cấp 1:
- Chọn **Styles Pane** ➜ **Tìm List Bullet** ➜ **Chỉnh sửa format** ➜ **Chọn add to template** ➜ **Nhấn OK**
- 📍 Đối với bullet cấp 2:
- Chọn **Styles Pane** ➜ **Tìm List Bullet 2** ➜ **Chọn Style type: Table** ➜ **Chỉnh sửa format** ➜ **Chọn add to template** ➜ **Nhấn OK**
- 📍 Đối với bảng:
- Chọn **Styles Pane** ➜ Chọn **New Style** ➜ **Chọn Style type: Table** ➜ **Chỉnh sửa format** ➜ Đổi tên thành `"New Table"` ➜ **Chọn add to template** ➜ **Nhấn OK**

    """)
    st.markdown("---")
    st.success("Ứng dụng được phát triển bởi VPI.")
    
st.subheader("1. Nhập thông tin đầu vào")

transcript_file = st.file_uploader("1. Tải lên file transcript (.docx)", type=["docx"])

st.subheader("2. Lựa chọn Template")
template_option = st.selectbox(
    "Bạn muốn sử dụng loại template nào?",
    ("Template VPI", "Template tùy chỉnh"),
    help="Chọn 'Template VPI' để dùng mẫu có sẵn hoặc 'Template tùy chỉnh' để tải lên file của riêng bạn."
)
template_file = None
if template_option == "Template tùy chỉnh":
    template_file = st.file_uploader("Tải lên file template .docx của bạn", type=["docx"])

st.subheader("3. Thông tin cơ bản")
# (MỚI) Chỉ hiện khi chọn Template tùy chỉnh
if template_option == "Template tùy chỉnh":
    st.info(
        "🔔 **Lưu ý đối với Template tùy chỉnh**\n\n"
        "- File template **bắt buộc** phải có đúng và đủ các biến sau, **đúng chính tả, không kèm mô tả `{# ... #}`**:\n"
        "  `{{TenCuocHop}}`, `{{ThoiGianCuocHop}}`, `{{DiaDiemCuocHop}}`, `{{TenChuTri}}`, `{{TenThuKy}}`.\n"
        "- Ví dụ **không hợp lệ**: `{{TenCuocHop}}{# ... #}` (không được kèm phần mô tả)."
    )
else:
    st.caption("Các trường bắt buộc đã có sẵn trong Template VPI.")
col1, col2 = st.columns(2)
with col1:
    meeting_name      = st.text_input("Tên cuộc họp")
    meeting_time      = st.text_input("Thời gian cuộc họp (VD: 10/9/2025)")
    meeting_location  = st.text_input("Địa điểm cuộc họp")
with col2:
    meeting_chair     = st.text_input("Tên chủ trì")
    meeting_secretary = st.text_input("Tên thư ký")

recipient_email = st.text_input("4. Email nhận kết quả của bạn")

# Nút chạy
if st.button("🚀 Tạo biên bản", type="primary"):
    # Đường dẫn template mặc định (nếu dùng Template VPI)
    default_path = "2025.VPI_BB hop 2025 1.docx"

    # 1) Kiểm tra bắt buộc (thiếu file/trường) -> báo đỏ + không chạy
    if not validate_inputs(
        template_option=template_option,
        transcript_file=transcript_file,
        template_file=template_file,
        meeting_name=meeting_name,
        meeting_time=meeting_time,
        meeting_location=meeting_location,
        meeting_chair=meeting_chair,
        meeting_secretary=meeting_secretary,
        recipient_email=recipient_email,
        default_template_path=default_path
    ):
        st.stop()  # CHẶN CHẠY TIẾP

    # 2) Xác định template để dùng (đã qua validate)
    template_to_use = None
    if template_option == "Template VPI":
        template_to_use = default_path
    else:
        template_to_use = template_file

    with st.spinner("⏳ Hệ thống đang xử lý..."):
        try:
            st.info("1/4 - Đang đọc và phân tích transcript...")
            doc = Document(transcript_file)
            transcript_content = "\n".join([para.text for para in doc.paragraphs])

            st.info("2/4 - Đang trích placeholders từ template...")
            placeholders = extract_vars_and_desc(template_to_use)

            # 2.1) Kiểm tra template có đủ placeholders bắt buộc không
            missing_ph = [k for k in REQUIRED_PLACEHOLDERS if k not in placeholders and k not in []]
            # Lưu ý: extract_vars_and_desc() chỉ trả về các biến có KÈM mô tả {#...#}.
            # Với 5 biến cơ bản yêu cầu "không kèm mô tả", ta vẫn chấp nhận vì phần điền thủ công override sau.
            # Tuy nhiên, nếu muốn ép buộc chặt chẽ hơn với Template tùy chỉnh, có thể đọc raw XML hoặc tự kiểm tra thêm.
            # Ở đây chỉ cảnh báo nếu hoàn toàn không thấy các biến này đâu trong template (cả có mô tả hay không).
            # Để kiểm tra "không kèm mô tả", ta sẽ kiểm sau khi mở Document(template_to_use) và scan text:
            # (Đoạn dưới làm kiểm tra mềm - cảnh báo nếu thiếu hẳn biến ở template.)

            try:
                tdoc = Document(template_to_use)
                ttext = "\n".join([p.text for p in tdoc.paragraphs])
                for ph in REQUIRED_PLACEHOLDERS:
                    if f"{{{{{ph}}}}}" not in ttext:
                        if ph not in missing_ph:
                            missing_ph.append(ph)
            except Exception:
                pass

            if missing_ph and template_option == "Template tùy chỉnh":
                st.error("❌ **Template tùy chỉnh thiếu các biến bắt buộc**: " + ", ".join(missing_ph) +
                         ".\nVui lòng cập nhật template rồi chạy lại.")
                st.stop()

            st.info("3/4 - Đang gọi AI để trích xuất nội dung...")
            llm_result = call_gemini_model(transcript_content, placeholders)

            if llm_result:
                # Ghi đè bằng input tay (trường bắt buộc)
                manual_inputs = {
                    'TenCuocHop':       meeting_name,
                    'ThoiGianCuocHop':  meeting_time,
                    'DiaDiemCuocHop':   meeting_location,
                    'TenChuTri':        meeting_chair,
                    'TenThuKy':         meeting_secretary
                }
                llm_result.update(manual_inputs)

                st.info("4/4 - Đang tạo file biên bản Word...")
                docx_buffer = fill_template_to_buffer(template_to_use, llm_result)
                if docx_buffer:
                    st.success("✅ Tạo biên bản thành công!")
                    st.download_button(
                        "⬇️ Tải về biên bản",
                        data=docx_buffer,
                        file_name="Bienbancuochop.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    # Gửi email (nếu có)
                    if recipient_email:
                        if send_email_with_attachment(recipient_email, docx_buffer, filename="Bien_ban_cuoc_hop.docx"):
                            st.success("✉️ Đã gửi biên bản tới email của bạn.")
                else:
                    st.error("Không thể tạo file Word. Vui lòng kiểm tra lại file template.")
            else:
                st.error("Không thể lấy kết quả từ AI. Vui lòng thử lại.")
        except Exception as e:
            st.error(f"Đã xảy ra lỗi: {e}")
